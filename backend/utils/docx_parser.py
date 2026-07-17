"""
DOCX Parser - Extracts text from DOCX files using python-docx.
"""

from typing import Optional


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from a DOCX file preserving structure."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Add style info as hints for section detection
                if paragraph.style and paragraph.style.name:
                    style = paragraph.style.name.lower()
                    if 'heading' in style:
                        text_parts.append(f"\n{text}")
                    else:
                        text_parts.append(text)
                else:
                    text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts) if text_parts else None

    except Exception as e:
        print(f"[DOCX Parser] Failed: {e}")
        return None
